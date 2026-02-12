#!/usr/bin/env python3
"""Fix three issues in main-anon.docx for JoEL submission:
1. Set double-spaced, 12pt Times New Roman
2. Add figure/table caption numbering
3. Convert footnotes to endnotes (LAST -- XML surgery that python-docx would overwrite)
"""
import zipfile
import os
import xml.etree.ElementTree as ET
from copy import deepcopy

WML = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"

SRC = "main-anon.docx"
DST = "English_kinship_terms_taboo_to_syntax_anon.docx"

# === PHASE 1: python-docx operations (formatting + captions) ===

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

# First, regenerate from pandoc to have a clean starting point
os.system(
    "pandoc main-anon.tex --from latex --to docx --citeproc "
    "--bibliography=references.bib --csl=unified-linguistics.csl "
    f"-o {SRC} 2>&1"
)
print("Regenerated clean docx from tex")

doc = Document(SRC)

# --- Fix 1: Double-spacing, 12pt Times New Roman (skip headings) ---
HEADING_STYLES = {"Title", "Subtitle", "Abstract Title",
                  "Heading 1", "Heading 2", "Heading 3", "Heading 4"}

for p in doc.paragraphs:
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    for run in p.runs:
        run.font.name = "Times New Roman"
        if p.style.name not in HEADING_STYLES:
            run.font.size = Pt(12)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)

# Set default body style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

# Set heading styles: Times New Roman, bold, appropriate sizes
from docx.shared import RGBColor
for sname, sz, bold, italic in [
    ("Title", 16, True, False),
    ("Abstract Title", 12, True, False),
    ("Heading 1", 14, True, False),
    ("Heading 2", 12, True, True),
    ("Heading 3", 12, True, False),
]:
    try:
        s = doc.styles[sname]
        s.font.name = "Times New Roman"
        s.font.size = Pt(sz)
        s.font.bold = bold
        s.font.italic = italic
        s.font.color.rgb = RGBColor(0, 0, 0)
        s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    except KeyError:
        pass

# Clear run-level overrides on headings so they inherit style (keep TNR)
for p in doc.paragraphs:
    if p.style.name in HEADING_STYLES:
        for run in p.runs:
            run.font.size = None
            run.font.bold = None
            run.font.italic = None
            run.font.color.rgb = None

# Add first-line indent to body paragraphs (not headings, captions, quotes, lists)
from docx.shared import Inches
NO_INDENT = HEADING_STYLES | {"Table Caption", "Image Caption", "Captioned Figure",
                               "Block Text", "List Number", "List Bullet",
                               "First Paragraph", "Author", "Date"}
first_after_heading = False
for p in doc.paragraphs:
    sn = p.style.name
    if sn in HEADING_STYLES:
        first_after_heading = True
        continue
    if sn not in NO_INDENT and p.text.strip():
        if first_after_heading:
            # First paragraph after a heading: no indent (standard style)
            first_after_heading = False
        else:
            p.paragraph_format.first_line_indent = Inches(0.5)
    elif not p.text.strip():
        pass  # blank lines don't reset the flag
    else:
        first_after_heading = False

print("Fix 1 done: double-spaced, 12pt Times New Roman")

# --- Fix 2: Caption numbering ---
fig_num = 0
tab_num = 0

for p in doc.paragraphs:
    text = p.text.strip()

    # Table captions: pandoc uses "Table Caption" style
    if p.style.name == "Table Caption" and not text.startswith("Table "):
        tab_num += 1
        if p.runs:
            p.runs[0].text = f"Table {tab_num}: " + p.runs[0].text
        print(f"  Table {tab_num}: {text[:60]}")

    # Figure captions: pandoc uses "Image Caption" style
    if p.style.name == "Image Caption" and not text.startswith("Figure "):
        fig_num += 1
        if p.runs:
            p.runs[0].text = f"Figure {fig_num}: " + p.runs[0].text
        print(f"  Figure {fig_num}: {text[:60]}")

print(f"Fix 2 done: numbered {tab_num} tables and {fig_num} figures")

# --- Fix 2b: Strip borders from example tables ---
import re as _re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

example_tables = 0
for table in doc.tables:
    first_cell = table.rows[0].cells[0].text.strip()
    if _re.match(r'\(\d+\)', first_cell):
        # This is an example table — remove all borders
        tbl = table._tbl
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        existing = tblPr.find(qn('w:tblBorders'))
        if existing is not None:
            tblPr.remove(existing)
        borders = OxmlElement('w:tblBorders')
        for bname in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{bname}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            borders.append(b)
        tblPr.append(borders)
        # Also remove cell-level borders
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.find(qn('w:tcPr'))
                if tcPr is not None:
                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)
        # Set narrow first column, wide second column
        table.columns[0].width = Inches(0.6)
        table.columns[1].width = Inches(5.4)
        example_tables += 1

print(f"Fix 2b done: stripped borders from {example_tables} example tables")

# Save after python-docx operations
doc.save(DST)
print("Saved after python-docx fixes")

# === PHASE 2: XML surgery for footnotes → endnotes ===
# Must happen AFTER python-docx save, since python-docx would overwrite our XML

# Register namespaces
for prefix, uri in {
    "w": WML, "r": REL,
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "mc": "http://schemas.openxmlformats.org/markup-compatibility/2006",
    "w14": "http://schemas.microsoft.com/office/word/2010/wordml",
    "w15": "http://schemas.microsoft.com/office/word/2012/wordml",
    "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
    "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
    "o": "urn:schemas-microsoft-com:office:office",
    "v": "urn:schemas-microsoft-com:vml",
}.items():
    ET.register_namespace(prefix, uri)

# Read all zip contents
with zipfile.ZipFile(DST, "r") as zin:
    names = zin.namelist()
    files = {name: zin.read(name) for name in names}

# Parse document.xml and footnotes.xml
doc_xml = ET.fromstring(files["word/document.xml"])
fn_xml = ET.fromstring(files["word/footnotes.xml"])

# Create endnotes.xml with separator entries
en_xml = ET.Element(f"{{{WML}}}endnotes")
for eid, etype in [("-1", "continuationSeparator"), ("0", "separator")]:
    en = ET.SubElement(en_xml, f"{{{WML}}}endnote")
    en.set(f"{{{WML}}}id", eid)
    en.set(f"{{{WML}}}type", etype)
    p = ET.SubElement(en, f"{{{WML}}}p")
    r = ET.SubElement(p, f"{{{WML}}}r")
    if etype == "separator":
        ET.SubElement(r, f"{{{WML}}}separator")
    else:
        ET.SubElement(r, f"{{{WML}}}continuationSeparator")

# Move real footnotes (id not 0 or -1) to endnotes
moved = 0
for fn_elem in list(fn_xml):
    fn_id = fn_elem.get(f"{{{WML}}}id")
    if fn_id and fn_id not in ("0", "-1"):
        en_elem = deepcopy(fn_elem)
        en_elem.tag = f"{{{WML}}}endnote"
        # Fix internal refs and styles
        for ref in en_elem.iter(f"{{{WML}}}footnoteRef"):
            ref.tag = f"{{{WML}}}endnoteRef"
        for sty in en_elem.iter(f"{{{WML}}}rStyle"):
            v = sty.get(f"{{{WML}}}val", "")
            if "Footnote" in v:
                sty.set(f"{{{WML}}}val", v.replace("Footnote", "Endnote"))
        for sty in en_elem.iter(f"{{{WML}}}pStyle"):
            v = sty.get(f"{{{WML}}}val", "")
            if "Footnote" in v:
                sty.set(f"{{{WML}}}val", v.replace("Footnote", "Endnote"))
        en_xml.append(en_elem)
        fn_xml.remove(fn_elem)
        moved += 1

# In document.xml: footnoteReference → endnoteReference
for fnref in doc_xml.iter(f"{{{WML}}}footnoteReference"):
    fnref.tag = f"{{{WML}}}endnoteReference"
for sty in doc_xml.iter(f"{{{WML}}}rStyle"):
    v = sty.get(f"{{{WML}}}val", "")
    if "Footnote" in v:
        sty.set(f"{{{WML}}}val", v.replace("Footnote", "Endnote"))

print(f"Fix 3: moved {moved} footnotes → endnotes")

# Set endnote numbering to arabic (decimal) in settings.xml
if "word/settings.xml" in files:
    settings_xml = ET.fromstring(files["word/settings.xml"])
    # Find or create endnotePr
    enpr = settings_xml.find(f"{{{WML}}}endnotePr")
    if enpr is None:
        enpr = ET.SubElement(settings_xml, f"{{{WML}}}endnotePr")
    # Set numFmt to decimal
    numfmt = enpr.find(f"{{{WML}}}numFmt")
    if numfmt is None:
        numfmt = ET.SubElement(enpr, f"{{{WML}}}numFmt")
    numfmt.set(f"{{{WML}}}val", "decimal")
    files["word/settings.xml"] = ET.tostring(settings_xml, xml_declaration=True, encoding="UTF-8")
    print("  Set endnote numbering to arabic")

# Ensure EndnoteReference style has superscript in styles.xml
if "word/styles.xml" in files:
    styles_xml = ET.fromstring(files["word/styles.xml"])
    # Find or create EndnoteReference style
    en_ref_style = None
    for s in styles_xml.iter(f"{{{WML}}}style"):
        if s.get(f"{{{WML}}}styleId") == "EndnoteReference":
            en_ref_style = s
            break
    if en_ref_style is None:
        # Create the style
        en_ref_style = ET.SubElement(styles_xml, f"{{{WML}}}style")
        en_ref_style.set(f"{{{WML}}}type", "character")
        en_ref_style.set(f"{{{WML}}}styleId", "EndnoteReference")
        name_el = ET.SubElement(en_ref_style, f"{{{WML}}}name")
        name_el.set(f"{{{WML}}}val", "endnote reference")
    # Ensure rPr with superscript
    rpr = en_ref_style.find(f"{{{WML}}}rPr")
    if rpr is None:
        rpr = ET.SubElement(en_ref_style, f"{{{WML}}}rPr")
    vertAlign = rpr.find(f"{{{WML}}}vertAlign")
    if vertAlign is None:
        vertAlign = ET.SubElement(rpr, f"{{{WML}}}vertAlign")
    vertAlign.set(f"{{{WML}}}val", "superscript")
    files["word/styles.xml"] = ET.tostring(styles_xml, xml_declaration=True, encoding="UTF-8")
    print("  Set EndnoteReference style to superscript")

# Update file contents
files["word/document.xml"] = ET.tostring(doc_xml, xml_declaration=True, encoding="UTF-8")
files["word/footnotes.xml"] = ET.tostring(fn_xml, xml_declaration=True, encoding="UTF-8")
files["word/endnotes.xml"] = ET.tostring(en_xml, xml_declaration=True, encoding="UTF-8")

# Ensure endnotes.xml in [Content_Types].xml
ct_xml = ET.fromstring(files["[Content_Types].xml"])
if not any(e.get("PartName") == "/word/endnotes.xml" for e in ct_xml):
    ov = ET.SubElement(ct_xml, "Override")
    ov.set("PartName", "/word/endnotes.xml")
    ov.set("ContentType",
           "application/vnd.openxmlformats-officedocument.wordprocessingml.endnotes+xml")
files["[Content_Types].xml"] = ET.tostring(ct_xml, xml_declaration=True, encoding="UTF-8")

# Ensure endnotes relationship in document.xml.rels
rels_key = "word/_rels/document.xml.rels"
if rels_key in files:
    rels_xml = ET.fromstring(files[rels_key])
    if not any(r.get("Target") == "endnotes.xml" for r in rels_xml):
        existing = {r.get("Id") for r in rels_xml}
        n = 1
        while f"rId{n}" in existing:
            n += 1
        new_rel = ET.SubElement(rels_xml, "Relationship")
        new_rel.set("Id", f"rId{n}")
        new_rel.set("Type",
                     "http://schemas.openxmlformats.org/officeDocument/2006/relationships/endnotes")
        new_rel.set("Target", "endnotes.xml")
    files[rels_key] = ET.tostring(rels_xml, xml_declaration=True, encoding="UTF-8")

# Write final zip
all_names = set(names) | {"word/endnotes.xml"}
with zipfile.ZipFile(DST, "w", zipfile.ZIP_DEFLATED) as zout:
    for name in all_names:
        if name in files:
            zout.writestr(name, files[name])

print(f"Fix 3 done: endnotes in {DST}")
print("\nAll fixes applied.")
